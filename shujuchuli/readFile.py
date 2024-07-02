from docx import Document
# import csv
import os.path
import pandas as pd
import xlrd


# import office


# 文本文件读取
# file1 = open(
#     'D:\\浏览器下载\\ai\\邗江区政务门牌事项梳理（114）\\邗江区政务门牌事项梳理（114）\\邗江文体旅局（3）\\体育类校外培训机构审批（个人、法人）\\体育类校外培训机构审批网办攻略.docx',
#     'r')
# print(file)

# with open(
#         'D:\\浏览器下载\\ai\\邗江区政务门牌事项梳理（114）\\邗江区政务门牌事项梳理（114）\\邗江文体旅局（3）\\体育类校外培训机构审批（个人、法人）\\体育类校外培训机构审批.xls',
#         'r') as file1:
#     content = file.read()
#     print(content)

# 读取文件夹 并输出内部文件
# path = 'D:\\浏览器下载\\ai\\邗江区政务门牌事项梳理（114）\\邗江区政务门牌事项梳理（114）\\邗江文体旅局（3）\\体育类校外培训机构审批（个人、法人）'  # 打开桌面位置处的文件夹hello，注意格式\\


# result = os.listdir(path)

# 读取文本文件
def read_docx(file_path):
    docx = Document(file_path)
    for paragraph in docx.paragraphs:
        return paragraph.text


# 读取xls
def read_xls(file_path):
    workbook = xlrd.open_workbook(file_path)
    for i in range(0, 2):
        if workbook.sheet_by_index(i) is not None:
            sheet = workbook.sheet_by_index(i)
            for row_idx in range(sheet.nrows):
                for col_idx in range(sheet.ncols):
                    cell_value = sheet.cell_value(row_idx, col_idx)
                    return cell_value


def traverse_directory(directory):
    for root, dirs, files in os.walk(directory):
        for name in files:
            file_path = os.path.join(root, name)
            if ".docx" in name:
                print(1)
                file.write(str(read_docx(file_path)))
            elif ".xlsx" in name:
                print(2)
                file.write(str(pd.read_excel(file_path)))
            elif ".xls" in name:
                file.write(str(read_xls(file_path)))
        for name in dirs:
            traverse_directory(os.path.join(root, name))

    # if __name__ == '__main__':
    # doc转docx


command = '\\邗江文体旅局（3）\\娱乐场所从事娱乐场所经营活动审批（娱乐场所经营许可）（法人）'
path = 'D:\\浏览器下载\\ai\\邗江区政务门牌事项梳理（114）\\邗江区政务门牌事项梳理（114）' + command  # 打开桌面位置处的文件夹hello，注意格式\\

with open('./file/娱乐场所从事娱乐场所经营活动审批（娱乐场所经营许可）（法人）.txt', 'a+', encoding="utf-8") as file:
    # office.word.doc2docx(path, path)
    traverse_directory(path)
# for temp in traverse_directory(path):
#     print(temp)
# if ".doc" in temp:
# temp.replace('\\', "\\")
# read_doc(temp)

# for file in result:
#     # print(file)
#     if ".doc" in file:
#         read_doc(path + "\\" + file)

# 文档文件读取
# doc = Document(
#     "D:\\浏览器下载\\ai\\邗江区政务门牌事项梳理（114）\\邗江区政务门牌事项梳理（114）\\邗江文体旅局（3）\\体育类校外培训机构审批（个人、法人）\\体育类校外培训机构审批网办攻略.docx")
# for paragraph in doc.paragraphs:
#     print(paragraph.text)

# 需要用到的库：python-docx、pandas、xlrd、python-office
